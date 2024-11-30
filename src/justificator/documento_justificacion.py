from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import xml.etree.ElementTree as ET
import os

from utilidades import imprimir_error_y_salir, imprimir_error, imprimir_correcto, imprimir_warning

class DocumentoJustificacion:
    def __init__(self, file_name, doc=None):
        self.doc = doc if doc else Document()
        self.file_name = file_name
        self.contador_claves_texto_a_texto_totales = 0
        self.contador_clave_texto_a_texto_encontrada = 0
        self.contador_reemplazos_texto_a_texto_totales = 0
        self.contador_claves_texto_a_imagenes_totales = 0
        self.contador_clave_texto_a_imagenes_encontrada = 0
        self.contador_imagenes_added_total = 0
        self.contador_claves_enlaces_totales = 0
        self.contador_clave_enlace_encontrada = 0
        self.contador_enlaces_added_total = 0

    def cargar_documento(self, template_folder="justificacion_template"):
        doc_path = os.path.join(template_folder, self.file_name)
        try:
            self.doc = Document(doc_path)
            imprimir_correcto(f"Documento {self.file_name} cargado correctamente desde {doc_path}")
        except Exception as e:
            imprimir_error_y_salir(f"Error al cargar el documento {doc_path}: {e}")

    def guardar_documento(self, output_folder="justificacion"):
        save_path = os.path.join(output_folder, self.file_name)
        try:
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)

            self.doc.save(save_path)
            imprimir_correcto(f"Documento guardado correctamente en {save_path}")
        except Exception as e:
            imprimir_error_y_salir(f"Error al guardar el documento {save_path}: {e}")

    def reemplazar_textos(self, textos_a_reemplazar, frontmatter, autoagregar_enlaces=False):
        self.contador_claves_texto_a_texto_totales += len(textos_a_reemplazar)
        if not frontmatter:
            return False

        # Crear diccionario solo con los textos que existen en frontmatter
        diccionario_reemplazo = {texto: frontmatter[texto] for texto in textos_a_reemplazar if frontmatter.get(texto)}

        for texto_original, texto_reemplazado in diccionario_reemplazo.items():
            texto_fue_encontrado = False
            es_un_enlace = autoagregar_enlaces and texto_reemplazado.startswith("http")

            if es_un_enlace:
                self.contador_claves_enlaces_totales += 1

            for parrafo in self.doc.paragraphs:
                for run in parrafo.runs:
                    texto_a_buscar = f"___{texto_original}___"
                    if texto_a_buscar in run.text:
                        texto_fue_encontrado = True
                        self.contador_reemplazos_texto_a_texto_totales += 1

                        if es_un_enlace:
                            run.text = run.text.replace(texto_a_buscar, "")
                            enlace_exito = self.agregar_hipervinculo(parrafo, texto_reemplazado)
                            if enlace_exito:
                                self.contador_enlaces_added_total += 1
                        else:
                            run.text = run.text.replace(texto_a_buscar, texto_reemplazado)

            # Actualizar contadores
            if texto_fue_encontrado:
                self.contador_clave_texto_a_texto_encontrada += 1
                if es_un_enlace:
                    self.contador_clave_enlace_encontrada += 1

    def reemplazar_imagenes(self, imagenes_a_reemplazar, width_inches:float=5, height_inches:float=3):
        self.contador_claves_texto_a_imagenes_totales += len(imagenes_a_reemplazar)

        if imagenes_a_reemplazar:
            for clave in imagenes_a_reemplazar:
                self.substituir_texto_por_imagenes(clave, imagenes_a_reemplazar[clave], width_inches, height_inches)

    def substituir_texto_por_imagenes(self, texto, imagenes, width_inches, height_inches):
        try:
            if not imagenes:
                imprimir_error("Imagen no proporcionada para el reemplazo.")
                return False

            if not texto:
                imprimir_error("No hay ningún texto que reemplazar.")
                return False

            for parrafo in self.doc.paragraphs:
                if f"___img_{texto}___" in parrafo.text:
                    parrafo.clear()
                    run = parrafo.add_run()
                    for imagen in imagenes:
                        run.add_picture(imagen, width=Inches(width_inches), height=Inches(height_inches))
                        run.add_break()
                        #run.add_break()
                        self.contador_imagenes_added_total += 1
                    self.contador_clave_texto_a_imagenes_encontrada += 1
                    return True

        except Exception as e:
            imprimir_error(f"Error al reemplazar la imagen: {e}")
            return False

    def obtener_dimenciones_de_imagen(self, parrafo_number):
        """ Legacy Method used when I used to manually replace images in parrragraphs """
        parrafo = self.doc.paragraphs[parrafo_number]
        if not self.parrafo_es_imagen(parrafo):
            imprimir_error("El parrafo no tiene imagen")
            return False

        # Parsear el XML del párrafo para obtener los valores de ancho y alto
        xml_content = parrafo._p.xml
        tree = ET.ElementTree(ET.fromstring(xml_content))
        root = tree.getroot()
        extent = root.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')

        if extent is not None:
            width = int(extent.attrib['cx'])  # El ancho en EMU
            height = int(extent.attrib['cy'])  # El alto en EMU

            # Convertir EMU a pulgadas para mantener el tamaño en el nuevo reemplazo
            inches_width = width / 914400
            inches_height = height / 914400
            return inches_width, inches_height
        else:
            imprimir_error("No se encontraron dimensiones de imagen en el párrafo.")
            return False

    def parrafo_es_imagen(self, parrafo):
        return 'graphicData' in parrafo._p.xml

    def agregar_hipervinculo(self, parrafo, url):
        """
        Inserta un hipervínculo en un párrafo sin modificar el texto existente.
        Aplica el formato de hipervínculo (color azul y subrayado).
        """
        try:
            # Crear el nodo de enlace
            part = self.doc.part
            HYPERLINK_RELATIONSHIP = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'
            r_id = part.relate_to(url, HYPERLINK_RELATIONSHIP, is_external=True)

            # Crear el elemento de hipervínculo (w:hyperlink)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            # Crear el run para el texto del enlace
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')  # Propiedades del run (formato)

            # Agregar subrayado (w:u) - formato típico de hipervínculo
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')  # Subrayado sencillo
            rPr.append(u)

            # Agregar color azul (w:color)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Azul
            rPr.append(color)

            # Agregar las propiedades de formato al run
            new_run.append(rPr)

            # Crear el texto del enlace
            texto_elemento = OxmlElement('w:t')
            texto_elemento.text = url
            new_run.append(texto_elemento)

            # Añadir el run al hipervínculo
            hyperlink.append(new_run)

            # Añadir el hipervínculo al párrafo
            parrafo._p.append(hyperlink)
            return True
        except Exception as e:
            imprimir_error(f"Error añadiendo enlace: {e}")
            return False

    def imprimir_estadisticas(self):

        mensaje_inicial = f"------- Resultados del documento: {self.file_name} -------"
        mensaje_textos = f"Textos a textos buscados y encontrados: {self.contador_clave_texto_a_texto_encontrada}/{self.contador_claves_texto_a_texto_totales} (Total Reemplazos Textos: {self.contador_reemplazos_texto_a_texto_totales})"
        mensaje_imagenes = f"Textos a imagenes buscados y encontrados: {self.contador_clave_texto_a_imagenes_encontrada}/{self.contador_claves_texto_a_imagenes_totales} (Total Imagenes Added: {self.contador_imagenes_added_total})"
        mensaje_enlaces = f"enlaces agregados: {self.contador_clave_enlace_encontrada}/{self.contador_claves_enlaces_totales} (Total Enlaces: {self.contador_enlaces_added_total})"

        textos_correcto = self.contador_clave_texto_a_texto_encontrada == self.contador_claves_texto_a_texto_totales
        imagenes_correctas = self.contador_clave_texto_a_imagenes_encontrada == self.contador_claves_texto_a_imagenes_totales
        enlaces_correctos = self.contador_clave_enlace_encontrada == self.contador_claves_enlaces_totales

        print()
        if textos_correcto and imagenes_correctas and enlaces_correctos:
            imprimir_correcto(mensaje_inicial)
        elif textos_correcto or imagenes_correctas or enlaces_correctos:
            imprimir_warning(mensaje_inicial)
        else:
            imprimir_error(mensaje_inicial)

        imprimir_correcto(mensaje_textos) if textos_correcto else imprimir_error(mensaje_textos)
        imprimir_correcto(mensaje_imagenes) if imagenes_correctas else imprimir_error(mensaje_imagenes)
        imprimir_correcto(mensaje_enlaces) if enlaces_correctos else imprimir_error(mensaje_enlaces)

        if textos_correcto and imagenes_correctas and enlaces_correctos:
            imprimir_correcto("------------------------------------------")
        elif textos_correcto or imagenes_correctas or enlaces_correctos:
            imprimir_warning("------------------------------------------")
        else:
            imprimir_error("------------------------------------------")

        print()


    def imprimir_estructura(self):
        for n, parrafo in enumerate(self.doc.paragraphs):
            print("-------")

            for nn, run in enumerate(parrafo.runs):
                print(f"{n}-{nn} : {run.text} ")

            print("<IMAGE>" if self.parrafo_es_imagen(parrafo) else "")
